import pyautogui
import time
import pyttsx3
import glob
import os
import re
from docx import Document
from docx.shared import Cm

num = input("Please select subject:\n 1.Math\n 2.Reasoning\n 3.English\n")
subject = {'1':'math', '2':'reasoning', '3':'english'}
	
def latest_file():
	path = r'C:\Users\HIMANSHU\Desktop\quant_screen\\'+subject[num]+'\\images\\*'
	total_files = glob.glob(path)
	latest_file = max(total_files, key=os.path.getctime)
	_, filenumber = os.path.split(latest_file)
	return filenumber

def guess_chapter():
	filenumber = latest_file()
	filenumber = re.findall(r'[a-zA-Z]+', filenumber[:-4])
	response = input(f'Want to continue with {filenumber[0]}, enter y else enter chapter name:\n')
	if response.lower() == 'y':
		filename = filenumber[0]
	else:
		create = input("If you want to create new doc, enter y:\n")
		if create.lower() =='y':
			filename = response.lower()
			doc = Document()
			doc.save(r'C:\Users\HIMANSHU\Desktop\quant_screen\\'+subject[num]+'\\'+filename+r'.docx')
		else:
			filename = response.lower()
			print(f"Opening {filename}.doc in edit mode.")
	return filename
	

def latest_number():
	#todo if folder new then no previous files
	filenumber = latest_file()
	if filename in filenumber:
		nums = re.findall(r'[0-9]+', filenumber)
	else:
		print(f"Image count starting from 1 as chapter name {filename} is not same as latest modified image {filenumber}.")
		nums = [0]
	return int(nums[0])+1

filename = guess_chapter()		
i=latest_number()

print(f"latest number {i}, latest chapter {filename}, subject {subject[num]}.")
def speak(bolo):
	engine = pyttsx3.init()
	engine.say(bolo)
	engine.runAndWait()
	engine.stop()
	
while(i>0):
 input("press enter to continue....")
 comment = input("comment:")
 for p in range(0,2):
  print("sleeping for "+str(p))
  time.sleep(1)
 pyautogui.screenshot(r'C:\Users\HIMANSHU\Desktop\quant_screen\\'+subject[num]+'\\images\\'+filename+str(i)+r'.png')
 document = Document(r'C:\Users\HIMANSHU\Desktop\quant_screen\\'+subject[num]+'\\'+filename+r'.docx')
 document.add_picture(r'C:\Users\HIMANSHU\Desktop\quant_screen\\'+subject[num]+'\\images\\'+filename+str(i)+r'.png', width=Cm(15.92))
 document.add_paragraph().add_run(comment).bold = True
 document.save(r'C:\Users\HIMANSHU\Desktop\quant_screen\\'+subject[num]+'\\'+filename+r'.docx')
 print('copied '+filename+str(i)+r'.png')
 speak("Completed")
 i+=1
 
 
#TO-DO
#Implement hand gesture controlled screenshot taking and take comment input from mike
#If fist then no comment
#If open hand then open cmd and take comment
#If two fingers then record video and attach as comment to screenshot
#Study opencv and numpy
#Continuously video will be camera will be ON and when it matches the previous stored fist,full or two fingers then it should trigger next action
